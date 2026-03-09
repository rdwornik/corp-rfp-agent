"""Answer inserter for Word documents.

Ported from src/rfp_answer_word.py -- inserts blue BY Response paragraphs.
"""

from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from corp_rfp_agent.agents.word.models import AnswerableBlock


# BY blue color: #0066CC
BY_BLUE_HEX = '0066CC'


def insert_answer_after(doc, paragraph_index: int, answer_text: str):
    """Insert a blue-colored BY Response paragraph after the given paragraph index.

    Exact port of insert_answer_after() from src/rfp_answer_word.py.
    Creates two runs:
    - "BY Response: " in bold blue (10pt)
    - Answer text in regular blue (10pt)
    """
    target_para = doc.paragraphs[paragraph_index]

    # Create new paragraph element after target
    new_para_elem = OxmlElement('w:p')
    target_para._element.addnext(new_para_elem)

    # "BY Response: " in bold blue
    run1 = OxmlElement('w:r')
    rpr1 = OxmlElement('w:rPr')
    bold1 = OxmlElement('w:b')
    rpr1.append(bold1)
    color1 = OxmlElement('w:color')
    color1.set(qn('w:val'), BY_BLUE_HEX)
    rpr1.append(color1)
    sz1 = OxmlElement('w:sz')
    sz1.set(qn('w:val'), '20')  # 10pt = 20 half-points
    rpr1.append(sz1)
    run1.append(rpr1)
    text1 = OxmlElement('w:t')
    text1.text = "BY Response: "
    text1.set(qn('xml:space'), 'preserve')
    run1.append(text1)
    new_para_elem.append(run1)

    # Answer text in regular blue
    run2 = OxmlElement('w:r')
    rpr2 = OxmlElement('w:rPr')
    color2 = OxmlElement('w:color')
    color2.set(qn('w:val'), BY_BLUE_HEX)
    rpr2.append(color2)
    sz2 = OxmlElement('w:sz')
    sz2.set(qn('w:val'), '20')
    rpr2.append(sz2)
    run2.append(rpr2)
    text2 = OxmlElement('w:t')
    text2.text = answer_text
    text2.set(qn('xml:space'), 'preserve')
    run2.append(text2)
    new_para_elem.append(run2)


def insert_blank_after(doc, paragraph_index: int):
    """Insert a blank paragraph after the given paragraph index."""
    target_para = doc.paragraphs[paragraph_index]
    new_para_elem = OxmlElement('w:p')
    target_para._element.addnext(new_para_elem)


def has_existing_response(doc, paragraph_index: int) -> bool:
    """Check if there's already a BY Response near this paragraph.

    Checks the paragraph itself and the immediately following paragraph
    for "BY Response:" text. We only check 2 positions to avoid
    false positives from responses belonging to adjacent sections.
    """
    paragraphs = doc.paragraphs
    for offset in range(0, 2):
        idx = paragraph_index + offset
        if idx >= len(paragraphs):
            break
        text = paragraphs[idx].text.strip()
        if "BY Response:" in text:
            return True
    return False
