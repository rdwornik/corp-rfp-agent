"""
rfp_answer_word.py
Answer client requirements in a Word document using KB + LLM.

Parses the document into a section tree using auto-detected headings
(bold + numbering patterns), then answers each leaf section with KB + LLM.

Usage:
  python src/rfp_answer_word.py --input "path/to/file.docx" --family network
  python src/rfp_answer_word.py --input "path/to/file.docx" --family network --dry-run
  python src/rfp_answer_word.py --input "path/to/file.docx" --family network --interactive
  python src/rfp_answer_word.py --input "path/to/file.docx" --family network \
    --answer-model gemini
"""

import argparse
import json
import os
import sys
import re
import time
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional

from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Project root setup
PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT / "src"))

from dotenv import load_dotenv
load_dotenv(PROJECT_ROOT / ".env", override=True)

import chromadb
from chromadb.utils import embedding_functions
from google import genai
from google.genai import types

from llm_router import MODELS, clean_bold_markdown, retry_with_backoff

# --- CONFIGURATION ---
FAMILY_CONFIG_PATH = PROJECT_ROOT / "data/kb/schema/family_config.json"
KB_JSON_PATH = PROJECT_ROOT / "data/kb/canonical/RFP_Database_UNIFIED_CANONICAL.json"
DB_PATH = PROJECT_ROOT / "data/kb/chroma_store"
COLLECTION_NAME = "rfp_knowledge_base"

BY_BLUE = RGBColor(0x00, 0x66, 0xCC)


# --- DATA CLASSES ---

@dataclass
class Section:
    level: int                    # 0 = chapter, 1 = section, 2 = subsection
    title: str
    para_index: int               # paragraph index in doc
    content_paragraphs: list = field(default_factory=list)   # [(index, text), ...]
    children: list = field(default_factory=list)             # list[Section]

    @property
    def full_content(self) -> str:
        """All content in this section (excluding children)."""
        return "\n".join(text for _, text in self.content_paragraphs)

    @property
    def full_content_with_children(self) -> str:
        """All content including children -- for context."""
        parts = [self.full_content]
        for child in self.children:
            parts.append(f"\n### {child.title}\n{child.full_content}")
        return "\n".join(parts)


@dataclass
class AnswerableBlock:
    section: Section
    breadcrumb: str              # "Architecture > Integration > SAP ECC6 Interfaces"
    content: str                 # all text in this section
    insert_after_para: int       # where to insert the BY response
    answer: str = ""
    kb_matches: list = field(default_factory=list)


# --- LOAD FAMILY CONFIG ---

def load_family_config():
    with open(FAMILY_CONFIG_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def get_family_display_name(family_code: str) -> str:
    config = load_family_config()
    family = config.get("families", {}).get(family_code, {})
    return family.get("display_name", family_code)


# --- HEADING DETECTION ---

def detect_heading_level(paragraph) -> Optional[int]:
    """Detect heading level from formatting patterns. Returns 0-3 or None."""
    text = paragraph.text.strip()
    if not text:
        return None

    style = paragraph.style.name if paragraph.style else "Normal"

    # Method 1: Proper Word heading styles (if they exist)
    if "Heading" in style:
        m = re.search(r'(\d+)', style)
        if m:
            return int(m.group(1)) - 1  # 0-indexed

    # Method 2: Auto-detect from bold + numbering patterns
    # Check if ALL runs with text are bold
    runs_with_text = [r for r in paragraph.runs if r.text.strip()]
    if not runs_with_text:
        return None
    all_bold = all(r.bold for r in runs_with_text)

    if not all_bold:
        return None  # Not a heading

    # Bold text -- now determine LEVEL from numbering pattern

    # Level 0: Roman numeral chapter (VIII. Architecture)
    if re.match(r'^[IVXLCDM]+[\.\)]\s', text):
        return 0

    # Level 1: Numbered section (1. Macro Technical Architecture)
    if re.match(r'^\d+[\.\)]\s', text):
        return 1

    # Level 2: Sub-numbered (1.1 or 1.1. pattern)
    if re.match(r'^\d+\.\d+[\.\)]?\s', text):
        return 2

    # Level 2: Short bold text without numbers = subsection header
    # (e.g. "Containerization", "Database", "APIs")
    if len(text) < 60:
        return 2

    return None  # Not a heading -- it's content


# --- SECTION TREE BUILDING ---

def build_section_tree(doc) -> list:
    """Parse document into section tree using auto-detected headings."""
    sections = []
    current_stack = []  # stack of sections by level

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        level = detect_heading_level(para)

        if level is not None:
            # This is a heading -- create new section
            section = Section(
                level=level,
                title=text,
                para_index=i,
            )

            # Find parent: pop stack until we find a section with lower level
            while current_stack and current_stack[-1].level >= level:
                current_stack.pop()

            if current_stack:
                current_stack[-1].children.append(section)
            else:
                sections.append(section)

            current_stack.append(section)
        else:
            # This is content -- add to deepest current section
            if current_stack:
                current_stack[-1].content_paragraphs.append((i, text))

    return sections


# --- COLLECT ANSWERABLE SECTIONS ---

def collect_answerable_sections(sections: list) -> list:
    """Collect leaf sections (and parents with own content) that need BY answers."""
    blocks = []

    def walk(section, parent_context=""):
        context = f"{parent_context} > {section.title}" if parent_context else section.title

        if section.children:
            # Section has children -- answer each child, not the parent
            # BUT if parent has its own content paragraphs, answer those too
            if section.content_paragraphs:
                blocks.append(AnswerableBlock(
                    section=section,
                    breadcrumb=context,
                    content=section.full_content,
                    insert_after_para=section.content_paragraphs[-1][0],
                ))
            for child in section.children:
                walk(child, context)
        else:
            # Leaf section -- answer it
            if section.content_paragraphs:
                blocks.append(AnswerableBlock(
                    section=section,
                    breadcrumb=context,
                    content=section.full_content,
                    insert_after_para=section.content_paragraphs[-1][0],
                ))

    for s in sections:
        walk(s)

    return blocks


# --- TREE DISPLAY ---

def print_section_tree(sections: list, indent: int = 0):
    """Print the section tree for confirmation."""
    for i, section in enumerate(sections):
        is_last = (i == len(sections) - 1)
        prefix = " " * indent
        connector = "`-- " if is_last else "|-- "

        content_count = len(section.content_paragraphs)
        children_count = len(section.children)

        info_parts = []
        if children_count:
            info_parts.append(f"{children_count} subsections")
        if content_count:
            info_parts.append(f"{content_count} paragraphs")
        info = f" ({', '.join(info_parts)})" if info_parts else ""

        print(f"{prefix}{connector}{section.title}{info}")

        if section.children:
            child_indent = indent + (4 if is_last else 4)
            print_section_tree(section.children, child_indent)


def count_sections_recursive(sections: list) -> int:
    """Count total sections in the tree."""
    count = 0
    for s in sections:
        count += 1
        count += count_sections_recursive(s.children)
    return count


# --- KB RETRIEVAL ---

class KBRetriever:
    """Lightweight KB retriever for Word RFP answering."""

    def __init__(self):
        print("[INFO] Connecting to ChromaDB...")
        self.client_db = chromadb.PersistentClient(path=str(DB_PATH))
        self.ef = embedding_functions.SentenceTransformerEmbeddingFunction(
            model_name="BAAI/bge-large-en-v1.5"
        )
        self.collection = self.client_db.get_collection(
            name=COLLECTION_NAME,
            embedding_function=self.ef
        )
        print(f"[OK] ChromaDB connected ({self.collection.count()} entries)")

        # Load KB for full answer lookup
        with open(KB_JSON_PATH, "r", encoding="utf-8") as f:
            raw_data = json.load(f)
        self.kb_lookup = {}
        for item in raw_data:
            kb_id = item.get("kb_id")
            domain = item.get("domain", "")
            if kb_id:
                self.kb_lookup[kb_id] = item
                if domain and not kb_id.startswith(f"{domain}_"):
                    self.kb_lookup[f"{domain}_{kb_id}"] = item

    def query(self, text: str, family_code: str = None, k: int = 5) -> list:
        """Query ChromaDB, optionally filtered by family/domain. Returns list of (item, distance)."""
        results = self.collection.query(query_texts=[text], n_results=k)

        matches = []
        if results["ids"] and results["ids"][0]:
            ids = results["ids"][0]
            distances = results["distances"][0] if "distances" in results else [1.0] * len(ids)
            for chroma_id, dist in zip(ids, distances):
                item = self.kb_lookup.get(chroma_id)
                if item:
                    matches.append((item, dist))

        # If family_code specified, boost matches from that domain
        if family_code and matches:
            family_matches = [(item, dist) for item, dist in matches
                              if item.get("domain") == family_code
                              or item.get("family_code") == family_code]
            platform_matches = [(item, dist) for item, dist in matches
                                if item.get("scope") == "platform"
                                or item.get("domain") == "platform"]
            other_matches = [(item, dist) for item, dist in matches
                             if (item, dist) not in family_matches
                             and (item, dist) not in platform_matches]
            matches = family_matches + platform_matches + other_matches

        return matches[:k]


# --- ANSWER GENERATION ---

def generate_answer(block: AnswerableBlock, family_display_name: str, model: str) -> str:
    """Generate a BY answer for a section block using KB context + LLM."""
    # Format KB entries
    if block.kb_matches:
        kb_parts = []
        for item, dist in block.kb_matches:
            similarity = 1.0 - dist  # ChromaDB returns L2 distance
            kb_parts.append(
                f"[KB {item.get('kb_id', '?')} | similarity: {similarity:.2f}]\n"
                f"Q: {item.get('canonical_question', '')}\n"
                f"A: {item.get('canonical_answer', '')}"
            )
        kb_context = "\n\n".join(kb_parts)
    else:
        kb_context = "(No relevant KB entries found. Answer from general Blue Yonder knowledge.)"

    prompt = f"""You are a Blue Yonder pre-sales engineer responding to a client RFP for {family_display_name}.

SECTION: {block.breadcrumb}

CLIENT REQUIREMENT:
{block.content}

RELEVANT KB ENTRIES:
{kb_context}

Write a professional response from Blue Yonder's perspective. Rules:
- Start with a DIRECT answer: "Blue Yonder [does/supports/provides] ..."
- Be specific about product capabilities
- Reference actual BY product names and features where relevant
- If KB entries provide relevant info, use that information
- If the requirement asks for something BY doesn't support, say so honestly: "This specific requirement would need to be addressed during implementation scoping."
- Keep it concise but complete: 2-5 sentences typically
- Do NOT make up features. If unsure, say "Blue Yonder can discuss this in detail during a technical deep-dive session."
- Write in English

Return ONLY the answer text, no JSON wrapping, no prefix."""

    model_config = MODELS.get(model, MODELS["gemini"])
    model_name = model_config["name"]

    def call_llm():
        client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))
        response = client.models.generate_content(
            model=model_name,
            contents=[{"role": "user", "parts": [{"text": prompt}]}],
            config=types.GenerateContentConfig(temperature=0.3, max_output_tokens=2048)
        )
        return response.text.strip() if response.text else ""

    raw = retry_with_backoff(call_llm)
    return clean_bold_markdown(raw)


# --- WORD DOC INSERTION ---

def insert_answer_after(doc, paragraph_index: int, answer_text: str):
    """Insert a blue-colored BY Response paragraph after the given paragraph index."""
    target_para = doc.paragraphs[paragraph_index]

    # Create new paragraph element after target
    new_para_elem = OxmlElement('w:p')
    target_para._element.addnext(new_para_elem)

    # "BY Response: " in bold blue
    run1 = OxmlElement('w:r')
    rpr1 = OxmlElement('w:rPr')
    bold1 = OxmlElement('w:b')
    rpr1.append(bold1)
    color1 = OxmlElement('w:color')
    color1.set(qn('w:val'), '0066CC')
    rpr1.append(color1)
    sz1 = OxmlElement('w:sz')
    sz1.set(qn('w:val'), '20')  # 10pt = 20 half-points
    rpr1.append(sz1)
    run1.append(rpr1)
    text1 = OxmlElement('w:t')
    text1.text = "BY Response: "
    text1.set(qn('xml:space'), 'preserve')
    run1.append(text1)
    new_para_elem.append(run1)

    # Answer text in regular blue
    run2 = OxmlElement('w:r')
    rpr2 = OxmlElement('w:rPr')
    color2 = OxmlElement('w:color')
    color2.set(qn('w:val'), '0066CC')
    rpr2.append(color2)
    sz2 = OxmlElement('w:sz')
    sz2.set(qn('w:val'), '20')
    rpr2.append(sz2)
    run2.append(rpr2)
    text2 = OxmlElement('w:t')
    text2.text = answer_text
    text2.set(qn('xml:space'), 'preserve')
    run2.append(text2)
    new_para_elem.append(run2)


def insert_blank_after(doc, paragraph_index: int):
    """Insert a blank paragraph after the given paragraph index."""
    target_para = doc.paragraphs[paragraph_index]
    new_para_elem = OxmlElement('w:p')
    target_para._element.addnext(new_para_elem)


# --- DRY RUN ---

def print_dry_run(blocks: list):
    """Print answerable sections without generating answers."""
    print(f"\n{'='*60}")
    print(f" DRY RUN -- {len(blocks)} answerable sections")
    print(f"{'='*60}\n")
    for i, block in enumerate(blocks, 1):
        content_preview = block.content[:120].replace("\n", " | ")
        print(f"  [{i:3d}] {block.breadcrumb}")
        print(f"        {content_preview}...")
        print()


# --- INTERACTIVE REVIEW ---

def interactive_review(blocks: list) -> list:
    """Let user review/edit answers interactively. Returns filtered blocks."""
    accepted = []
    total = len(blocks)

    for i, block in enumerate(blocks):
        print(f"\n{'='*60}")
        print(f" Section {i+1}/{total} | {block.breadcrumb}")
        print(f"{'='*60}")
        print(f"\n SECTION CONTENT:")
        for line in block.content.split("\n"):
            print(f"   {line}")

        print(f"\n BY RESPONSE: {block.answer}")

        if block.kb_matches:
            match_strs = []
            for item, dist in block.kb_matches[:3]:
                sim = 1.0 - dist
                match_strs.append(f"{item.get('kb_id', '?')} ({sim:.2f})")
            print(f"\n KB: {', '.join(match_strs)}")

        print(f"\n [Y] Accept  [E] Edit  [N] Skip  [A] Accept all remaining  [Q] Quit")
        choice = input("> ").strip().upper()

        if choice == "A":
            accepted.append(block)
            accepted.extend(blocks[i+1:])
            print(f"   Accepted all {len(blocks) - i} remaining.")
            break
        elif choice == "Q":
            print("   Quitting review.")
            break
        elif choice == "N":
            print("   Skipped.")
            continue
        elif choice == "E":
            print("   Enter new answer (end with empty line):")
            lines = []
            while True:
                line = input()
                if not line:
                    break
                lines.append(line)
            if lines:
                block.answer = " ".join(lines)
            accepted.append(block)
        else:  # Y or anything else
            accepted.append(block)

    return accepted


# --- MAIN ---

def main():
    parser = argparse.ArgumentParser(description="Answer client requirements in a Word document")
    parser.add_argument("--input", "-i", required=True, help="Path to input .docx file")
    parser.add_argument("--family", "-f", required=True, help="Product family code (e.g. network, wms, planning)")
    parser.add_argument("--answer-model", default="gemini",
                        help="LLM for answer generation (default: gemini = gemini-3.1-pro-preview)")
    parser.add_argument("--model", "-m", default=None,
                        help="Alias for --answer-model")
    parser.add_argument("--dry-run", action="store_true", help="Show detected sections without answering")
    parser.add_argument("--interactive", action="store_true", help="Review each answer before inserting")
    parser.add_argument("--kb-threshold", type=float, default=0.75, help="Min similarity for KB match (default: 0.75)")
    parser.add_argument("--top-k", type=int, default=5, help="Number of KB entries to retrieve (default: 5)")
    parser.add_argument("--yes", "-y", action="store_true", help="Skip tree confirmation prompt")
    args = parser.parse_args()

    # --model is alias for --answer-model
    if args.model:
        args.answer_model = args.model

    answer_model = args.answer_model

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"[ERROR] File not found: {input_path}")
        sys.exit(1)

    family_display = get_family_display_name(args.family)
    print(f"\n[INFO] RFP Word Answerer (section-tree mode)")
    print(f"   Input:  {input_path.name}")
    print(f"   Family: {family_display} ({args.family})")
    print(f"   Model:  {answer_model} ({MODELS.get(answer_model, {}).get('name', '?')})")
    if args.dry_run:
        print(f"   Mode:   DRY RUN")
    elif args.interactive:
        print(f"   Mode:   INTERACTIVE")
    print()

    # Step 1: Parse Word doc
    print("[1/7] Parsing Word document...")
    doc = Document(str(input_path))
    total_paras = len(doc.paragraphs)
    non_empty = sum(1 for p in doc.paragraphs if p.text.strip())
    print(f"   {total_paras} total paragraphs ({non_empty} non-empty)")

    # Step 2: Auto-detect headings and build section tree
    print(f"\n[2/7] Building section tree (programmatic heading detection)...")
    sections = build_section_tree(doc)
    total_sections = count_sections_recursive(sections)
    print(f"   {total_sections} sections detected")

    if not sections:
        print("[ERROR] No sections detected. The document may not have recognizable headings.")
        sys.exit(1)

    # Step 3: Print tree summary, confirm structure
    print(f"\n[3/7] Document structure:\n")
    print_section_tree(sections)

    blocks = collect_answerable_sections(sections)
    print(f"\n   Answerable sections: {len(blocks)}")

    if not blocks:
        print("[INFO] No answerable sections found.")
        sys.exit(0)

    if not args.dry_run and not args.yes:
        confirm = input("\nProceed? [Y/n] ").strip().lower()
        if confirm and confirm != "y":
            print("Aborted.")
            return

    # Dry run: just show and exit
    if args.dry_run:
        print_dry_run(blocks)
        return

    # Step 4: Query KB for each section
    print(f"\n[4/7] Retrieving KB context...")
    retriever = KBRetriever()
    kb_used_count = 0

    for i, block in enumerate(blocks):
        query_text = f"{block.section.title}\n{block.content}"
        matches = retriever.query(query_text, family_code=args.family, k=args.top_k)

        # Filter by similarity threshold
        filtered = [(item, dist) for item, dist in matches if (1.0 - dist) >= args.kb_threshold]
        block.kb_matches = filtered if filtered else matches[:2]

        if filtered:
            kb_used_count += 1

        if (i + 1) % 10 == 0 or i == len(blocks) - 1:
            print(f"   {i+1}/{len(blocks)} sections queried")

    print(f"   KB matches (>={args.kb_threshold} similarity): {kb_used_count}/{len(blocks)}")

    # Step 5: Generate answers
    print(f"\n[5/7] Generating answers ({answer_model})...")
    for i, block in enumerate(blocks):
        title_preview = block.section.title[:60]
        print(f"   [{i+1}/{len(blocks)}] {title_preview}...")
        block.answer = generate_answer(block, family_display, answer_model)
        if (i + 1) < len(blocks):
            time.sleep(0.3)  # rate limit courtesy

    total_answerable = len(blocks)

    # Step 6: Interactive review
    if args.interactive:
        blocks = interactive_review(blocks)
        print(f"\n   Accepted {len(blocks)} answers")

    if not blocks:
        print("[INFO] No answers to insert.")
        return

    # Step 7: Insert answers into Word doc (BACKWARDS to avoid index shift)
    print(f"\n[7/7] Inserting answers into document...")
    blocks_sorted = sorted(blocks, key=lambda b: b.insert_after_para, reverse=True)

    for block in blocks_sorted:
        insert_blank_after(doc, block.insert_after_para)
        insert_answer_after(doc, block.insert_after_para, block.answer)

    # Save output
    output_name = input_path.stem + "_BY_RESPONSE" + input_path.suffix
    output_path = input_path.parent / output_name
    doc.save(str(output_path))

    # Summary
    print(f"\n{'='*60}")
    print(f" RFP Answer Complete: {input_path.name}")
    print(f"   Sections detected: {total_sections}")
    print(f"   Answerable: {total_answerable}")
    print(f"   Answered: {len(blocks)}")
    if args.interactive:
        print(f"   Skipped: {total_answerable - len(blocks)}")
    print(f"   KB matches used: {kb_used_count}")
    print(f"")
    print(f"   Output: {output_path.name}")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()
