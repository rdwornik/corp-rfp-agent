"""
Word Document Analyzer using local Ollama LLM.

This module extracts tables from .docx files and uses a local Ollama model
to analyze the structure and identify questions/answer columns.
"""

import json
import re
from pathlib import Path
from typing import Optional

from docx import Document

# Optional Ollama import with fallback
try:
    import ollama
    OLLAMA_AVAILABLE = True
except ImportError:
    OLLAMA_AVAILABLE = False


def extract_document_text(doc: Document, max_chars: int = 5000) -> str:
    """Extract text content from document paragraphs (non-table content)."""
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    full_text = "\n".join(paragraphs)
    if len(full_text) > max_chars:
        full_text = full_text[:max_chars] + "...[truncated]"

    return full_text


def extract_tables(docx_path: str) -> list:
    """
    Extract all tables from a .docx file.

    Returns list of table dicts with structure:
    {
        "index": int,
        "rows": int,
        "cols": int,
        "cells": [[{row, col, text}, ...], ...]
    }
    """
    doc = Document(docx_path)
    tables = []

    for i, table in enumerate(doc.tables):
        table_data = {
            "index": i,
            "rows": len(table.rows),
            "cols": len(table.columns) if table.columns else 0,
            "cells": []
        }

        for row_idx, row in enumerate(table.rows):
            row_cells = []
            for col_idx, cell in enumerate(row.cells):
                row_cells.append({
                    "row": row_idx,
                    "col": col_idx,
                    "text": cell.text.strip()[:200]  # Truncate for LLM
                })
            table_data["cells"].append(row_cells)

        tables.append(table_data)

    return tables


def format_tables_for_llm(tables: list, max_rows_per_table: int = 10) -> str:
    """Format tables into a readable string for LLM analysis."""
    output = []

    for table in tables:
        output.append(f"\n=== TABLE {table['index']} ({table['rows']} rows x {table['cols']} cols) ===")

        # Only show first N rows to keep context size reasonable
        rows_to_show = min(table['rows'], max_rows_per_table)

        for row_idx in range(rows_to_show):
            if row_idx < len(table['cells']):
                row_cells = table['cells'][row_idx]
                cell_texts = [c['text'][:100] for c in row_cells]  # Further truncate
                output.append(f"Row {row_idx}: {' | '.join(cell_texts)}")

        if table['rows'] > max_rows_per_table:
            output.append(f"... ({table['rows'] - max_rows_per_table} more rows)")

    return "\n".join(output)


def build_analysis_prompt(document_text: str, tables_formatted: str) -> str:
    """Build the prompt for Ollama document analysis."""
    return f"""Analyze this Word document and return a JSON structure identifying the questions and where to put answers.

DOCUMENT TEXT (non-table content):
{document_text}

TABLES:
{tables_formatted}

TASK:
1. Identify which tables contain RFP/security questions
2. For each question table, identify:
   - Which column contains the questions
   - Which column should contain the answers
   - Whether there's a header row
3. Extract all questions with their locations
4. Identify any client/company names that should be anonymized

Return ONLY valid JSON with this exact structure:
{{
  "document_structure": "table",
  "client_names_found": ["Company Name 1"],
  "tables": [
    {{
      "table_index": 0,
      "purpose": "security questions",
      "question_column": 1,
      "answer_column": 2,
      "has_header_row": true,
      "data_rows_start": 1
    }}
  ],
  "questions": [
    {{
      "id": "q1",
      "table_index": 0,
      "row_index": 1,
      "question_column": 1,
      "answer_column": 2,
      "original_text": "Does the system support SSO?",
      "anonymized_text": "Does the system support SSO?",
      "category": "authentication"
    }}
  ]
}}

RULES:
- document_structure: "table" if questions are in tables, "text" if in paragraphs, "mixed" if both
- question_column/answer_column are 0-indexed column numbers
- Anonymize company names by replacing with [CLIENT]
- Category should be one of: authentication, encryption, compliance, infrastructure, integration, data, security, general
- Include ALL questions found, not just the first few

Return ONLY the JSON, no other text."""


def analyze_with_ollama(
    document_text: str,
    tables: list,
    model: str = "qwen2.5:7b"
) -> dict:
    """
    Send document structure to Ollama for analysis.

    Args:
        document_text: Non-table text from the document
        tables: Extracted table structures
        model: Ollama model name (default: qwen2.5:7b)

    Returns:
        Analysis dict with document structure, questions, and locations
    """
    if not OLLAMA_AVAILABLE:
        raise ImportError("Ollama package not installed. Run: pip install ollama")

    tables_formatted = format_tables_for_llm(tables)
    prompt = build_analysis_prompt(document_text, tables_formatted)

    print(f"[INFO] Sending document to Ollama ({model}) for analysis...")

    response = ollama.chat(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        format="json"
    )

    content = response['message']['content']

    # Parse JSON response
    try:
        analysis = json.loads(content)
    except json.JSONDecodeError as e:
        # Try to extract JSON from response if there's extra text
        json_match = re.search(r'\{[\s\S]*\}', content)
        if json_match:
            analysis = json.loads(json_match.group())
        else:
            raise ValueError(f"Failed to parse Ollama response as JSON: {e}\nResponse: {content[:500]}")

    return analysis


def analyze_document(
    docx_path: str,
    local_llm: str = "qwen2.5:7b",
    output_path: Optional[str] = None
) -> dict:
    """
    Main entry point for document analysis.

    Args:
        docx_path: Path to .docx file
        local_llm: Ollama model to use for analysis
        output_path: Optional path to save analysis JSON

    Returns:
        Analysis dict containing:
        - document_structure
        - client_names_found
        - tables (with column mappings)
        - questions (with locations and anonymized versions)
    """
    docx_path = Path(docx_path)

    if not docx_path.exists():
        raise FileNotFoundError(f"Document not found: {docx_path}")

    if not docx_path.suffix.lower() == '.docx':
        raise ValueError(f"Expected .docx file, got: {docx_path.suffix}")

    print(f"[INFO] Analyzing document: {docx_path.name}")

    # Load document
    doc = Document(str(docx_path))

    # Extract content
    document_text = extract_document_text(doc)
    tables = extract_tables(str(docx_path))

    print(f"[INFO] Found {len(tables)} tables in document")

    # Analyze with Ollama
    analysis = analyze_with_ollama(document_text, tables, model=local_llm)

    # Add metadata
    analysis['source_file'] = str(docx_path)
    analysis['total_tables'] = len(tables)
    analysis['total_questions'] = len(analysis.get('questions', []))

    print(f"[SUCCESS] Analysis complete: {analysis['total_questions']} questions found")

    # Save analysis if output path provided
    if output_path:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(analysis, f, indent=2, ensure_ascii=False)

        print(f"[SUCCESS] Analysis saved to: {output_path}")

    return analysis


def print_analysis_summary(analysis: dict) -> None:
    """Print a human-readable summary of the analysis."""
    print("\n" + "=" * 60)
    print("DOCUMENT ANALYSIS SUMMARY")
    print("=" * 60)

    print(f"\nDocument Structure: {analysis.get('document_structure', 'unknown')}")
    print(f"Total Tables: {analysis.get('total_tables', 0)}")
    print(f"Total Questions: {analysis.get('total_questions', 0)}")

    client_names = analysis.get('client_names_found', [])
    if client_names:
        print(f"\nClient Names Found (will be anonymized):")
        for name in client_names:
            print(f"  - {name}")

    tables_info = analysis.get('tables', [])
    if tables_info:
        print(f"\nTable Mappings:")
        for t in tables_info:
            print(f"  Table {t.get('table_index')}: {t.get('purpose', 'unknown')}")
            print(f"    Question Column: {t.get('question_column')}")
            print(f"    Answer Column: {t.get('answer_column')}")
            print(f"    Header Row: {t.get('has_header_row')}")

    questions = analysis.get('questions', [])
    if questions:
        print(f"\nQuestions Preview (first 5):")
        for q in questions[:5]:
            print(f"  [{q.get('id')}] {q.get('anonymized_text', '')[:60]}...")
            print(f"       Category: {q.get('category', 'general')}")
            print(f"       Location: Table {q.get('table_index')}, Row {q.get('row_index')}")

    print("\n" + "=" * 60)


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Analyze Word document structure with Ollama")
    parser.add_argument("--input", required=True, help="Path to .docx file")
    parser.add_argument("--model", default="qwen2.5:7b", help="Ollama model (default: qwen2.5:7b)")
    parser.add_argument("--output", help="Optional output path for analysis JSON")

    args = parser.parse_args()

    analysis = analyze_document(
        docx_path=args.input,
        local_llm=args.model,
        output_path=args.output
    )

    print_analysis_summary(analysis)
