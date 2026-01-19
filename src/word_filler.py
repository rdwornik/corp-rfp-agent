"""
Word Document Filler.

This module fills generated answers back into the original Word document
based on the analysis JSON from word_analyzer.
"""

import json
from pathlib import Path
from typing import Optional
from copy import deepcopy

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def load_analysis(analysis_path: str) -> dict:
    """Load analysis JSON from file."""
    with open(analysis_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def load_answers(answers_path: str) -> dict:
    """
    Load answers JSON from file.

    Expected format:
    {
        "q1": {"answer": "...", "confidence": 0.95},
        "q2": {"answer": "...", "confidence": 0.88},
        ...
    }
    """
    with open(answers_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def fill_document(
    docx_path: str,
    analysis: dict,
    answers: dict,
    output_path: Optional[str] = None,
    preserve_formatting: bool = True
) -> str:
    """
    Fill answers into a Word document based on analysis.

    Args:
        docx_path: Path to original .docx file
        analysis: Analysis dict from word_analyzer
        answers: Dict mapping question IDs to answer dicts
        output_path: Output path for filled document (default: {original}_FILLED.docx)
        preserve_formatting: Try to preserve original cell formatting

    Returns:
        Path to the filled document
    """
    docx_path = Path(docx_path)

    if not docx_path.exists():
        raise FileNotFoundError(f"Document not found: {docx_path}")

    # Determine output path
    if output_path:
        output_path = Path(output_path)
    else:
        output_path = docx_path.parent / f"{docx_path.stem}_FILLED{docx_path.suffix}"

    print(f"[INFO] Loading document: {docx_path.name}")

    # Load document
    doc = Document(str(docx_path))

    # Track fill statistics
    filled_count = 0
    skipped_count = 0
    error_count = 0

    # Get questions from analysis
    questions = analysis.get('questions', [])

    print(f"[INFO] Filling {len(questions)} answer cells...")

    for question in questions:
        q_id = question.get('id')
        table_index = question.get('table_index')
        row_index = question.get('row_index')
        answer_column = question.get('answer_column')

        # Get answer for this question
        answer_data = answers.get(q_id)

        if not answer_data:
            print(f"[WARNING] No answer found for {q_id}")
            skipped_count += 1
            continue

        answer_text = answer_data.get('answer', '')

        if not answer_text or answer_text == "Not in KB":
            print(f"[SKIP] No KB answer for {q_id}")
            skipped_count += 1
            continue

        # Validate indices
        if table_index is None or row_index is None or answer_column is None:
            print(f"[ERROR] Missing location info for {q_id}")
            error_count += 1
            continue

        # Get the table
        if table_index >= len(doc.tables):
            print(f"[ERROR] Table index {table_index} out of range for {q_id}")
            error_count += 1
            continue

        table = doc.tables[table_index]

        # Get the row
        if row_index >= len(table.rows):
            print(f"[ERROR] Row index {row_index} out of range for {q_id}")
            error_count += 1
            continue

        row = table.rows[row_index]

        # Get the answer cell
        if answer_column >= len(row.cells):
            print(f"[ERROR] Column index {answer_column} out of range for {q_id}")
            error_count += 1
            continue

        cell = row.cells[answer_column]

        # Fill the answer
        try:
            # Clear existing content
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.text = ""

            # Set new content
            if cell.paragraphs:
                paragraph = cell.paragraphs[0]
                paragraph.text = answer_text

                # Apply default formatting if not preserving
                if not preserve_formatting:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Calibri'
            else:
                cell.text = answer_text

            filled_count += 1

        except Exception as e:
            print(f"[ERROR] Failed to fill {q_id}: {e}")
            error_count += 1

    # Save filled document
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    print(f"\n[SUCCESS] Document saved: {output_path}")
    print(f"[STATS] Filled: {filled_count} | Skipped: {skipped_count} | Errors: {error_count}")

    return str(output_path)


def fill_document_from_files(
    docx_path: str,
    analysis_path: str,
    answers_path: str,
    output_path: Optional[str] = None
) -> str:
    """
    Convenience function to fill document from file paths.

    Args:
        docx_path: Path to original .docx file
        analysis_path: Path to analysis JSON
        answers_path: Path to answers JSON
        output_path: Optional output path

    Returns:
        Path to the filled document
    """
    analysis = load_analysis(analysis_path)
    answers = load_answers(answers_path)

    return fill_document(
        docx_path=docx_path,
        analysis=analysis,
        answers=answers,
        output_path=output_path
    )


def deanonymize_answers(answers: dict, client_names: list) -> dict:
    """
    Replace [CLIENT] placeholders with actual client name in answers.

    Args:
        answers: Dict of answers with potential [CLIENT] placeholders
        client_names: List of client names found during analysis

    Returns:
        Answers with [CLIENT] replaced by first client name (if any)
    """
    if not client_names:
        return answers

    client_name = client_names[0]  # Use first found client name
    result = deepcopy(answers)

    for q_id, answer_data in result.items():
        if 'answer' in answer_data:
            answer_data['answer'] = answer_data['answer'].replace('[CLIENT]', client_name)

    return result


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Fill answers into Word document")
    parser.add_argument("--input", required=True, help="Path to original .docx file")
    parser.add_argument("--analysis", required=True, help="Path to analysis JSON")
    parser.add_argument("--answers", required=True, help="Path to answers JSON")
    parser.add_argument("--output", help="Output path (default: {input}_FILLED.docx)")

    args = parser.parse_args()

    output = fill_document_from_files(
        docx_path=args.input,
        analysis_path=args.analysis,
        answers_path=args.answers,
        output_path=args.output
    )

    print(f"\nFilled document: {output}")
