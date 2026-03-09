"""
create_fixtures.py
Generate golden-file test fixtures for acceptance tests.

Usage:
    python tests/create_fixtures.py

Creates:
    tests/fixtures/excel_golden.xlsx
    tests/fixtures/word_golden.docx
"""

from pathlib import Path
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIXTURES_DIR = Path(__file__).parent / "fixtures"


# =============================================================================
# EXCEL GOLDEN FIXTURE
# =============================================================================

def create_excel_golden():
    """Create golden Excel fixture exercising all edge cases."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Requirements"

    green_fill = PatternFill(start_color="FF00FF00", end_color="FF00FF00", fill_type="solid")
    bold_font = Font(bold=True)
    header_fill = PatternFill(start_color="FFD9E1F2", end_color="FFD9E1F2", fill_type="solid")

    # Row 1: Headers
    headers = {
        "A": "ID",
        "B": "Category",
        "C": "Requirement",
        "D": "Priority",
        "E": "Vendor Response",
        "F": "Notes",
    }
    for col_letter, header_text in headers.items():
        cell = ws[f"{col_letter}1"]
        cell.value = header_text
        cell.font = bold_font
        cell.fill = header_fill

    # Row 2: Green cell — needs answer
    ws["A2"].value = "REQ-001"
    ws["B2"].value = "Integration"
    ws["C2"].value = "The solution must provide REST API endpoints for data exchange."
    ws["D2"].value = "High"
    ws["E2"].value = None  # empty — to be filled
    ws["E2"].fill = green_fill

    # Row 3: Green cell — needs answer
    ws["A3"].value = "REQ-002"
    ws["B3"].value = "Security"
    ws["C3"].value = "The solution must support SSO via SAML 2.0 and OpenID Connect."
    ws["D3"].value = "High"
    ws["E3"].value = None
    ws["E3"].fill = green_fill

    # Row 4: Already answered — should NOT be touched
    ws["A4"].value = "REQ-003"
    ws["B4"].value = "Architecture"
    ws["C4"].value = "Describe the deployment architecture."
    ws["D4"].value = "Medium"
    ws["E4"].value = "Blue Yonder supports cloud-native SaaS deployment on Microsoft Azure."

    # Row 5: Green cell — needs answer
    ws["A5"].value = "REQ-004"
    ws["B5"].value = "Data"
    ws["C5"].value = "The solution must support data encryption at rest and in transit."
    ws["D5"].value = "High"
    ws["E5"].value = None
    ws["E5"].fill = green_fill

    # Row 6: Already answered — should NOT be touched
    ws["A6"].value = "REQ-005"
    ws["B6"].value = "Compliance"
    ws["C6"].value = "Does the platform comply with GDPR?"
    ws["D6"].value = "High"
    ws["E6"].value = "Yes, the platform is fully GDPR compliant with data residency options."

    # Row 7: Merged cell spanning C7:D7 — question in merged range
    ws.merge_cells("C7:D7")
    ws["A7"].value = "REQ-006"
    ws["B7"].value = "Performance"
    ws["C7"].value = "The system must handle 10,000 concurrent users with sub-second response times."
    ws["E7"].value = None
    ws["E7"].fill = green_fill

    # Row 8: Formula in adjacent cell F8
    ws["A8"].value = "REQ-007"
    ws["B8"].value = "Scalability"
    ws["C8"].value = "Describe horizontal scaling capabilities."
    ws["D8"].value = "Medium"
    ws["E8"].value = None
    ws["E8"].fill = green_fill
    ws["F8"].value = "=LEN(E8)"

    # Row 9: Special characters in question
    ws["A9"].value = "REQ-008"
    ws["B9"].value = "i18n"
    ws["C9"].value = 'Does the system support "quotes" & unicode: \u0142\u00f3\u015b\u0107?'
    ws["D9"].value = "Low"
    ws["E9"].value = None
    ws["E9"].fill = green_fill

    # Row 10: Category header (bold, no answer cell, not green)
    ws["A10"].value = None
    ws["B10"].value = None
    ws["C10"].value = "Security Requirements"
    ws["C10"].font = bold_font
    ws["D10"].value = None
    ws["E10"].value = None

    # Adjust column widths
    ws.column_dimensions["C"].width = 60
    ws.column_dimensions["E"].width = 50

    output_path = FIXTURES_DIR / "excel_golden.xlsx"
    wb.save(str(output_path))
    print(f"[OK] Created {output_path}")
    return output_path


# =============================================================================
# WORD GOLDEN FIXTURE
# =============================================================================

def _add_bold_paragraph(doc, text):
    """Add a paragraph where all runs are bold."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    return p


def _add_blue_response(doc, text):
    """Add a blue BY Response paragraph (simulating previous run output)."""
    p = doc.add_paragraph()

    # "BY Response: " bold blue
    run1 = p.add_run("BY Response: ")
    run1.bold = True
    run1.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    run1.font.size = Pt(10)

    # Answer text regular blue
    run2 = p.add_run(text)
    run2.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    run2.font.size = Pt(10)

    return p


def create_word_golden():
    """Create golden Word fixture exercising all edge cases."""
    doc = Document()

    # Chapter heading — level 0 (bold, roman numeral)
    _add_bold_paragraph(doc, "VIII. Architecture")

    # Section 1 — level 1 (bold, numbered)
    _add_bold_paragraph(doc, "1. Integration Capabilities")
    doc.add_paragraph(
        "The solution must integrate seamlessly with existing enterprise systems "
        "and support multiple integration patterns."
    )

    # Subsection: APIs — level 2
    _add_bold_paragraph(doc, "APIs")
    doc.add_paragraph(
        "The solution must expose RESTful APIs or GraphQL endpoints, "
        "secured using OAuth 2.0 or API key authentication."
    )
    doc.add_paragraph(
        "API rate limiting and throttling must be configurable per client."
    )

    # Subsection: File Exchanges — level 2
    _add_bold_paragraph(doc, "File Exchanges")
    doc.add_paragraph(
        "The solution must support file-based data exchange via SFTP."
    )
    doc.add_paragraph(
        "Supported formats must include CSV, XML, and JSON."
    )
    doc.add_paragraph(
        "File processing must support scheduling and error handling with retry logic."
    )

    # Subsection: Event Bus — level 2
    _add_bold_paragraph(doc, "Event Bus")
    doc.add_paragraph(
        "The solution must support event-driven integration via Apache Kafka "
        "or Azure Event Hub."
    )
    doc.add_paragraph(
        "Real-time synchronization of master data and transactional events is required."
    )

    # Section 2 — Security (has a table)
    _add_bold_paragraph(doc, "2. Security")
    doc.add_paragraph(
        "The solution must meet enterprise security standards as outlined below."
    )

    # Add a table with security requirements
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    # Header row
    table.cell(0, 0).text = "Requirement"
    table.cell(0, 1).text = "Standard"
    table.cell(0, 2).text = "Priority"
    # Data rows
    table.cell(1, 0).text = "Data encryption at rest"
    table.cell(1, 1).text = "AES-256"
    table.cell(1, 2).text = "Mandatory"
    table.cell(2, 0).text = "Data encryption in transit"
    table.cell(2, 1).text = "TLS 1.2+"
    table.cell(2, 2).text = "Mandatory"
    table.cell(3, 0).text = "SSO Integration"
    table.cell(3, 1).text = "SAML 2.0"
    table.cell(3, 2).text = "Mandatory"

    doc.add_paragraph(
        "All security controls must be auditable and compliant with ISO 27001."
    )

    # Section 3 — Monitoring (has existing blue BY response)
    _add_bold_paragraph(doc, "3. Monitoring")
    doc.add_paragraph(
        "The solution must provide comprehensive monitoring and alerting capabilities."
    )
    _add_blue_response(
        doc,
        "Blue Yonder provides built-in monitoring through Azure Monitor integration, "
        "with configurable alerts and dashboards for real-time visibility."
    )

    output_path = FIXTURES_DIR / "word_golden.docx"
    doc.save(str(output_path))
    print(f"[OK] Created {output_path}")
    return output_path


# =============================================================================
# MAIN
# =============================================================================

if __name__ == "__main__":
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    create_excel_golden()
    create_word_golden()
    print("\n[DONE] All fixtures created in tests/fixtures/")
