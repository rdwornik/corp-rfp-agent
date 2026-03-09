"""Tests for answer_inserter -- validates blue text insertion."""

import pytest
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

from corp_rfp_agent.agents.word.section_parser import (
    build_section_tree,
    collect_answerable_sections,
)
from corp_rfp_agent.agents.word.answer_inserter import (
    insert_answer_after,
    insert_blank_after,
    has_existing_response,
)

FIXTURES_DIR = Path(__file__).resolve().parents[1] / "fixtures"
MOCK_ANSWER = "Blue Yonder supports this capability through our cloud-native platform."


@pytest.fixture
def golden_doc():
    return Document(str(FIXTURES_DIR / "word_golden.docx"))


def test_blue_text_insertion_position(golden_doc, tmp_path):
    """BY Response inserted after the last paragraph of each section."""
    sections = build_section_tree(golden_doc)
    blocks = collect_answerable_sections(sections)

    apis_block = [b for b in blocks if "APIs" in b.breadcrumb][0]
    insert_idx = apis_block.insert_after_para
    insert_answer_after(golden_doc, insert_idx, MOCK_ANSWER)

    out_path = tmp_path / "test_output.docx"
    golden_doc.save(str(out_path))
    doc2 = Document(str(out_path))

    found = False
    for para in doc2.paragraphs:
        if "BY Response:" in para.text and MOCK_ANSWER in para.text:
            found = True
            break
    assert found, "BY Response not found in output document"


def test_blue_text_formatting(golden_doc, tmp_path):
    """Inserted text has blue color (0066CC) and 'BY Response:' prefix."""
    sections = build_section_tree(golden_doc)
    blocks = collect_answerable_sections(sections)

    block = blocks[0]
    insert_answer_after(golden_doc, block.insert_after_para, MOCK_ANSWER)

    out_path = tmp_path / "test_format.docx"
    golden_doc.save(str(out_path))
    doc2 = Document(str(out_path))

    for para in doc2.paragraphs:
        if "BY Response:" in para.text:
            runs = para.runs
            assert len(runs) >= 2, "Expected at least 2 runs (prefix + answer)"

            # First run: bold blue prefix
            prefix_run = runs[0]
            assert prefix_run.bold is True
            assert prefix_run.text.strip().startswith("BY Response")

            # Check color via XML
            for run in runs:
                rpr = run._element.find(qn('w:rPr'))
                if rpr is not None:
                    color_elem = rpr.find(qn('w:color'))
                    if color_elem is not None:
                        assert color_elem.get(qn('w:val')) == '0066CC'
            break
    else:
        pytest.fail("BY Response paragraph not found")


def test_table_preservation(golden_doc, tmp_path):
    """Tables in the document survive answer insertion intact."""
    tables_before = len(golden_doc.tables)
    assert tables_before >= 1

    table_content_before = []
    for table in golden_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_content_before.append(cell.text)

    sections = build_section_tree(golden_doc)
    blocks = collect_answerable_sections(sections)
    blocks_sorted = sorted(blocks, key=lambda b: b.insert_after_para, reverse=True)

    for block in blocks_sorted:
        insert_blank_after(golden_doc, block.insert_after_para)
        insert_answer_after(golden_doc, block.insert_after_para, MOCK_ANSWER)

    out_path = tmp_path / "test_tables.docx"
    golden_doc.save(str(out_path))
    doc2 = Document(str(out_path))

    assert len(doc2.tables) == tables_before

    table_content_after = []
    for table in doc2.tables:
        for row in table.rows:
            for cell in row.cells:
                table_content_after.append(cell.text)
    assert table_content_before == table_content_after


def test_document_integrity(golden_doc, tmp_path):
    """Output document opens and has valid structure after insertions."""
    para_count_before = len(golden_doc.paragraphs)

    sections = build_section_tree(golden_doc)
    blocks = collect_answerable_sections(sections)
    num_blocks = len(blocks)

    blocks_sorted = sorted(blocks, key=lambda b: b.insert_after_para, reverse=True)
    for block in blocks_sorted:
        insert_blank_after(golden_doc, block.insert_after_para)
        insert_answer_after(golden_doc, block.insert_after_para, MOCK_ANSWER)

    out_path = tmp_path / "test_integrity.docx"
    golden_doc.save(str(out_path))

    doc2 = Document(str(out_path))
    # Paragraph count increased by 2 per block (answer + blank)
    expected_increase = num_blocks * 2
    assert len(doc2.paragraphs) == para_count_before + expected_increase


def test_has_existing_response(golden_doc):
    """has_existing_response detects the existing BY Response in Monitoring section."""
    sections = build_section_tree(golden_doc)
    blocks = collect_answerable_sections(sections)

    monitoring_block = [b for b in blocks if "Monitoring" in b.breadcrumb][0]
    assert has_existing_response(golden_doc, monitoring_block.insert_after_para)

    apis_block = [b for b in blocks if "APIs" in b.breadcrumb][0]
    assert not has_existing_response(golden_doc, apis_block.insert_after_para)
