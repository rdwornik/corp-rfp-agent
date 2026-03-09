"""Tests for WordAgent v2 with mock LLM and KB."""

import pytest
from pathlib import Path
from docx import Document

from corp_rfp_agent.agents.word.agent import WordAgent
from corp_rfp_agent.agents.word.models import WordAgentResult
from corp_rfp_agent.anonymize.service import Anonymizer
from corp_rfp_agent.core.types import LLMResponse, KBMatch

FIXTURES_DIR = Path(__file__).resolve().parents[1] / "fixtures"


class MockLLMClient:
    """Mock LLM that returns a fixed answer."""
    def __init__(self, answer="Blue Yonder supports this capability."):
        self._answer = answer
        self.calls = []

    def generate(self, prompt, *, model=None, system_prompt=None, temperature=0.3, max_tokens=2000):
        self.calls.append(prompt)
        return LLMResponse(text=self._answer, model=model or "mock", provider="mock")

    def generate_json(self, prompt, *, model=None, system_prompt=None):
        return {}


class MockKBClient:
    """Mock KB that returns fixed matches."""
    def __init__(self, matches=None):
        self._matches = matches or []

    def query(self, question, *, family=None, category=None, top_k=5, threshold=0.75):
        return self._matches

    def count(self, family=None):
        return len(self._matches)

    def families(self):
        return {}


def test_process_dry_run():
    """Dry run returns answerable sections without generating answers."""
    agent = WordAgent(
        llm_client=MockLLMClient(),
        kb_client=MockKBClient(),
    )
    result = agent.process(
        FIXTURES_DIR / "word_golden.docx",
        dry_run=True,
    )
    assert result.answerable == 6
    assert result.answered == 0
    assert result.total_sections == 7


def test_process_writes_answers(tmp_path):
    """Process with mock LLM writes answers to document."""
    agent = WordAgent(
        llm_client=MockLLMClient(answer="Mock answer for testing."),
        kb_client=MockKBClient(),
    )
    output = tmp_path / "output.docx"
    result = agent.process(
        FIXTURES_DIR / "word_golden.docx",
        output_path=output,
    )
    # Monitoring already has a BY Response, so skip_existing filters it
    assert result.answered == 5
    assert result.errors == 0
    assert output.exists()

    # Verify answers written
    doc = Document(str(output))
    by_responses = [p for p in doc.paragraphs if "BY Response:" in p.text]
    # 5 new + 1 existing = 6
    assert len(by_responses) == 6


def test_process_skip_existing(tmp_path):
    """Existing BY Response in Monitoring section is not duplicated."""
    agent = WordAgent(
        llm_client=MockLLMClient(answer="New answer"),
        kb_client=MockKBClient(),
    )
    output = tmp_path / "output.docx"
    result = agent.process(
        FIXTURES_DIR / "word_golden.docx",
        output_path=output,
        skip_existing=True,
    )
    # Monitoring should be skipped (has existing response)
    assert result.answered == 5

    # Verify original monitoring response still present
    doc = Document(str(output))
    found_original = False
    for p in doc.paragraphs:
        if "Azure Monitor" in p.text:
            found_original = True
            break
    assert found_original


def test_process_applies_anonymization(tmp_path):
    """Anonymizer is called before LLM."""
    llm = MockLLMClient(answer="Answer for [Customer]")
    anonymizer = Anonymizer(client_name="AcmeCorp")
    agent = WordAgent(
        llm_client=llm,
        kb_client=MockKBClient(),
        anonymizer=anonymizer,
    )
    output = tmp_path / "output.docx"
    agent.process(FIXTURES_DIR / "word_golden.docx", output_path=output)

    # LLM was called for answerable sections (5 non-existing)
    assert len(llm.calls) == 5


def test_process_applies_overrides(tmp_path):
    """Overrides are applied to final answer."""
    import yaml
    overrides_path = tmp_path / "overrides.yaml"
    data = {"overrides": [
        {"id": "OVR-TEST", "find": "Mock answer", "replace": "Overridden answer"},
    ]}
    overrides_path.write_text(yaml.dump(data), encoding="utf-8")

    from corp_rfp_agent.overrides.store import YAMLOverrideStore
    store = YAMLOverrideStore(yaml_path=overrides_path)

    agent = WordAgent(
        llm_client=MockLLMClient(answer="Mock answer for testing."),
        kb_client=MockKBClient(),
        override_store=store,
    )
    output = tmp_path / "output.docx"
    result = agent.process(FIXTURES_DIR / "word_golden.docx", output_path=output)

    # Verify overrides applied in output
    doc = Document(str(output))
    found_override = False
    for p in doc.paragraphs:
        if "Overridden answer" in p.text:
            found_override = True
            break
    assert found_override


def test_result_tracks_counts(tmp_path):
    """WordAgentResult tracks correct counts."""
    agent = WordAgent(
        llm_client=MockLLMClient(answer="A"),
        kb_client=MockKBClient(),
    )
    output = tmp_path / "output.docx"
    result = agent.process(FIXTURES_DIR / "word_golden.docx", output_path=output)

    assert result.total_sections == 7
    assert result.answerable == 6
    assert result.answered == 5  # Monitoring skipped (existing response)
    assert result.skipped == 0
    assert result.errors == 0
    assert len(result.results) == 5
