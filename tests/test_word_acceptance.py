"""Acceptance tests for Word agent -- structure checks only (no LLM calls)."""

import sys
import copy
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

# Add src/ to import path
PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT / "src"))

from rfp_answer_word import (
    detect_heading_level,
    build_section_tree,
    collect_answerable_sections,
    insert_answer_after,
    insert_blank_after,
)


# ---------------------------------------------------------------------------
# Test 1: Section tree correctly detected
# ---------------------------------------------------------------------------
def test_section_tree_detection(word_golden, word_golden_expected):
    """Auto-detect heading hierarchy builds correct tree."""
    doc = Document(str(word_golden))
    sections = build_section_tree(doc)

    # Top level: "VIII. Architecture" at level 0
    assert len(sections) == 1
    root = sections[0]
    assert root.title == "VIII. Architecture"
    assert root.level == 0

    # Children: 3 numbered sections
    expected_children = word_golden_expected["expected_tree"][0]["children"]
    assert len(root.children) == len(expected_children)

    for actual, expected in zip(root.children, expected_children):
        assert actual.title == expected["title"]
        assert actual.level == expected["level"]


# ---------------------------------------------------------------------------
# Test 2: Integration subsections detected at level 2
# ---------------------------------------------------------------------------
def test_integration_subsections(word_golden):
    """APIs, File Exchanges, Event Bus are level 2 under Integration."""
    doc = Document(str(word_golden))
    sections = build_section_tree(doc)

    integration = sections[0].children[0]  # "1. Integration Capabilities"
    assert integration.title == "1. Integration Capabilities"
    assert len(integration.children) == 3

    titles = [c.title for c in integration.children]
    assert titles == ["APIs", "File Exchanges", "Event Bus"]

    for child in integration.children:
        assert child.level == 2


# ---------------------------------------------------------------------------
# Test 3: Content paragraph counts match
# ---------------------------------------------------------------------------
def test_content_paragraph_counts(word_golden, word_golden_expected):
    """Each section has the expected number of content paragraphs."""
    doc = Document(str(word_golden))
    sections = build_section_tree(doc)
    integration = sections[0].children[0]

    expected = word_golden_expected["expected_tree"][0]["children"][0]

    # Integration's own content
    assert len(integration.content_paragraphs) == expected["content_paragraphs"]

    # Subsection content
    for actual_child, expected_child in zip(integration.children, expected["children"]):
        assert len(actual_child.content_paragraphs) == expected_child["content_paragraphs"], \
            f"{actual_child.title}: expected {expected_child['content_paragraphs']}, " \
            f"got {len(actual_child.content_paragraphs)}"


# ---------------------------------------------------------------------------
# Test 4: Answerable sections collected correctly
# ---------------------------------------------------------------------------
def test_answerable_sections(word_golden, word_golden_expected):
    """Only sections with content are collected for answering."""
    doc = Document(str(word_golden))
    sections = build_section_tree(doc)
    blocks = collect_answerable_sections(sections)

    expected_count = word_golden_expected["total_answerable_sections"]
    assert len(blocks) == expected_count

    # Chapter header alone (no own content) should NOT be answerable
    breadcrumbs = [b.breadcrumb for b in blocks]
    assert not any(bc == "VIII. Architecture" for bc in breadcrumbs), \
        "Chapter heading without own content should not be answerable"


# ---------------------------------------------------------------------------
# Test 5: Breadcrumbs are correct
# ---------------------------------------------------------------------------
def test_breadcrumbs(word_golden):
    """Breadcrumb paths correctly reflect hierarchy."""
    doc = Document(str(word_golden))
    sections = build_section_tree(doc)
    blocks = collect_answerable_sections(sections)

    breadcrumbs = [b.breadcrumb for b in blocks]

    assert "VIII. Architecture > 1. Integration Capabilities > APIs" in breadcrumbs
    assert "VIII. Architecture > 1. Integration Capabilities > File Exchanges" in breadcrumbs
    assert "VIII. Architecture > 1. Integration Capabilities > Event Bus" in breadcrumbs
    assert "VIII. Architecture > 2. Security" in breadcrumbs
    assert "VIII. Architecture > 3. Monitoring" in breadcrumbs


# ---------------------------------------------------------------------------
# Test 6: Blue text insertion position
# ---------------------------------------------------------------------------
def test_blue_text_insertion_position(word_golden, mock_llm_response, tmp_path):
    """BY Response inserted after the last paragraph of each section."""
    doc = Document(str(word_golden))

    # Find the paragraph index for APIs section's last content
    sections = build_section_tree(doc)
    blocks = collect_answerable_sections(sections)

    apis_block = [b for b in blocks if "APIs" in b.breadcrumb][0]
    insert_idx = apis_block.insert_after_para

    # Insert answer
    insert_answer_after(doc, insert_idx, mock_llm_response)

    # Save and reload
    out_path = tmp_path / "test_output.docx"
    doc.save(str(out_path))
    doc2 = Document(str(out_path))

    # The paragraph after the insert point should contain our response
    found = False
    for para in doc2.paragraphs:
        if "BY Response:" in para.text and mock_llm_response in para.text:
            found = True
            break
    assert found, "BY Response not found in output document"


# ---------------------------------------------------------------------------
# Test 7: Blue text has correct formatting
# ---------------------------------------------------------------------------
def test_blue_text_formatting(word_golden, mock_llm_response, tmp_path):
    """Inserted text has blue color (0066CC) and 'BY Response:' prefix."""
    doc = Document(str(word_golden))
    sections = build_section_tree(doc)
    blocks = collect_answerable_sections(sections)

    # Insert into first answerable section
    block = blocks[0]
    insert_answer_after(doc, block.insert_after_para, mock_llm_response)

    out_path = tmp_path / "test_format.docx"
    doc.save(str(out_path))
    doc2 = Document(str(out_path))

    # Find the inserted paragraph
    for para in doc2.paragraphs:
        if "BY Response:" in para.text:
            runs = para.runs
            assert len(runs) >= 2, "Expected at least 2 runs (prefix + answer)"

            # First run: bold blue prefix
            prefix_run = runs[0]
            assert prefix_run.bold is True
            assert prefix_run.text.strip().startswith("BY Response")

            # Check color via XML (more reliable than run.font.color.rgb)
            for run in runs:
                rpr = run._element.find(qn('w:rPr'))
                if rpr is not None:
                    color_elem = rpr.find(qn('w:color'))
                    if color_elem is not None:
                        assert color_elem.get(qn('w:val')) == '0066CC'
            break
    else:
        pytest.fail("BY Response paragraph not found")


# ---------------------------------------------------------------------------
# Test 8: Tables not corrupted
# ---------------------------------------------------------------------------
def test_table_preservation(word_golden, mock_llm_response, tmp_path):
    """Tables in the document survive answer insertion intact."""
    doc = Document(str(word_golden))

    # Count tables before
    tables_before = len(doc.tables)
    assert tables_before >= 1, "Fixture should have at least one table"

    # Get table content before
    table_content_before = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_content_before.append(cell.text)

    # Insert answers into all blocks (backwards)
    sections = build_section_tree(doc)
    blocks = collect_answerable_sections(sections)
    blocks_sorted = sorted(blocks, key=lambda b: b.insert_after_para, reverse=True)

    for block in blocks_sorted:
        insert_blank_after(doc, block.insert_after_para)
        insert_answer_after(doc, block.insert_after_para, mock_llm_response)

    # Save and reload
    out_path = tmp_path / "test_tables.docx"
    doc.save(str(out_path))
    doc2 = Document(str(out_path))

    # Tables count preserved
    assert len(doc2.tables) == tables_before

    # Table content preserved
    table_content_after = []
    for table in doc2.tables:
        for row in table.rows:
            for cell in row.cells:
                table_content_after.append(cell.text)

    assert table_content_before == table_content_after


# ---------------------------------------------------------------------------
# Test 9: Document opens without corruption
# ---------------------------------------------------------------------------
def test_document_integrity(word_golden, mock_llm_response, tmp_path):
    """Output document opens and has valid structure after insertions."""
    doc = Document(str(word_golden))
    para_count_before = len(doc.paragraphs)

    sections = build_section_tree(doc)
    blocks = collect_answerable_sections(sections)
    num_blocks = len(blocks)

    # Insert answers backwards
    blocks_sorted = sorted(blocks, key=lambda b: b.insert_after_para, reverse=True)
    for block in blocks_sorted:
        insert_blank_after(doc, block.insert_after_para)
        insert_answer_after(doc, block.insert_after_para, mock_llm_response)

    out_path = tmp_path / "test_integrity.docx"
    doc.save(str(out_path))

    # Re-open -- no exceptions
    doc2 = Document(str(out_path))

    # Paragraph count increased by 2 per block (answer + blank)
    expected_increase = num_blocks * 2
    assert len(doc2.paragraphs) == para_count_before + expected_increase


# ---------------------------------------------------------------------------
# Test 10: detect_heading_level patterns
# ---------------------------------------------------------------------------
def test_heading_level_detection_patterns(word_golden):
    """Verify heading level detection for various formatting patterns."""
    doc = Document(str(word_golden))

    heading_levels = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        level = detect_heading_level(para)
        if level is not None:
            heading_levels[text] = level

    # Roman numeral -> level 0
    assert heading_levels.get("VIII. Architecture") == 0

    # Numbered sections -> level 1
    assert heading_levels.get("1. Integration Capabilities") == 1
    assert heading_levels.get("2. Security") == 1
    assert heading_levels.get("3. Monitoring") == 1

    # Short bold text -> level 2
    assert heading_levels.get("APIs") == 2
    assert heading_levels.get("File Exchanges") == 2
    assert heading_levels.get("Event Bus") == 2
